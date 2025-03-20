#!/usr/bin/env python3
"""
DOCX to Markdown converter

This script converts Microsoft Word documents to Markdown format.
"""

import argparse
import sys
from pathlib import Path
import re
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.table import _Cell, Table


class DOCX2Markdown:
    """Convert Microsoft Word documents to Markdown format."""

    def __init__(self, input_path, output_path=None):
        """
        Initialize the converter.

        Args:
            input_path (str): Path to the input DOCX file
            output_path (str, optional): Path to the output Markdown file
        """
        self.input_path = Path(input_path)
        self.output_path = Path(output_path) if output_path else None
        
        # Validate input file exists and is a DOCX file
        if not self.input_path.exists():
            raise FileNotFoundError(f"Input file not found: {self.input_path}")
        
        if self.input_path.suffix.lower() != '.docx':
            raise ValueError(f"Input file must be a DOCX file: {self.input_path}")

    def _extract_hyperlinks(self, paragraph):
        """
        Extract hyperlinks from paragraph and replace with markdown links.
        
        Args:
            paragraph: docx paragraph object
            
        Returns:
            str: Text with hyperlinks formatted as markdown
        """
        text = paragraph.text
        links = {}
        
        # Get all hyperlinks from paragraph
        for rel in paragraph.part.rels:
            if paragraph.part.rels[rel].reltype == RT.HYPERLINK:
                target = paragraph.part.rels[rel].target_ref
                for i, run in enumerate(paragraph.runs):
                    if run.element.xpath(".//w:hyperlink"):
                        if run.text in text:
                            links[run.text] = target
                        
        # Replace hyperlinks with markdown format
        for text_fragment, url in links.items():
            if text_fragment and url:
                text = text.replace(text_fragment, f"[{text_fragment}]({url})")
                
        return text

    def _process_paragraph(self, paragraph):
        """
        Convert a paragraph to markdown format.
        
        Args:
            paragraph: docx paragraph object
            
        Returns:
            str: Markdown formatted paragraph text
        """
        if not paragraph.text.strip():
            return ""
            
        # Check paragraph style and format accordingly
        style = paragraph.style.name.lower()
        text = self._extract_hyperlinks(paragraph)
        
        # Process text styles (bold, italic)
        runs_markdown = []
        for run in paragraph.runs:
            run_text = run.text
            
            # Apply formatting
            if run.bold and run.italic:
                run_text = f"***{run_text}***"
            elif run.bold:
                run_text = f"**{run_text}**"
            elif run.italic:
                run_text = f"*{run_text}*"
                
            runs_markdown.append(run_text)
            
        text = "".join(runs_markdown) if runs_markdown else text
        
        # Handle headings
        if style.startswith('heading'):
            level = 1
            if style[-1].isdigit():
                level = int(style[-1])
            heading_marker = '#' * level
            return f"{heading_marker} {text}"
            
        # Handle lists
        if paragraph.style.name.startswith('List'):
            if any(p.strip().startswith(('•', '*', '-', '•')) for p in text.split('\n')):
                # Bullet list
                if not text.startswith(('•', '*', '-', '•')):
                    text = f"* {text}"
            else:
                # Check if this might be part of a numbered list
                match = re.match(r'^(\d+\.)', text.strip())
                if match:
                    # Already has a number, ensure proper format
                    number = match.group(1)
                    text = f"{number} {text[len(number):].strip()}"
                else:
                    # Add a list marker
                    text = f"* {text}"
                    
        return text

    def _process_table(self, table):
        """
        Convert a table to markdown format.
        
        Args:
            table: docx table object
            
        Returns:
            str: Markdown formatted table
        """
        if not table.rows:
            return ""
            
        md_table = []
        
        # Process header row
        header = []
        for cell in table.rows[0].cells:
            header.append(cell.text.strip() or " ")
        md_table.append("| " + " | ".join(header) + " |")
        
        # Add separator row
        separator = []
        for _ in header:
            separator.append("---")
        md_table.append("| " + " | ".join(separator) + " |")
        
        # Process data rows
        for row in table.rows[1:]:
            row_cells = []
            for cell in row.cells:
                row_cells.append(cell.text.strip() or " ")
            md_table.append("| " + " | ".join(row_cells) + " |")
            
        return "\n".join(md_table)

    def convert(self):
        """
        Convert the DOCX file to Markdown format.
        
        Returns:
            str: Markdown representation of the DOCX document
        """
        try:
            # Load the DOCX document
            doc = Document(self.input_path)
            
            # Process document elements
            markdown_elements = []
            
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    for paragraph in doc.paragraphs:
                        if paragraph._element == element:
                            markdown_elements.append(self._process_paragraph(paragraph))
                            break
                elif element.tag.endswith('tbl'):  # Table
                    for table in doc.tables:
                        if table._element == element:
                            markdown_elements.append(self._process_table(table))
                            break
            
            # Join all elements with proper spacing
            markdown = "\n\n".join(filter(None, markdown_elements))
            
            # Write to output file if specified
            if self.output_path:
                with open(self.output_path, 'w', encoding='utf-8') as f:
                    f.write(markdown)
                print(f"Successfully converted {self.input_path} to {self.output_path}")
            
            return markdown
            
        except Exception as e:
            print(f"Error converting DOCX to Markdown: {e}", file=sys.stderr)
            raise


def main():
    """Parse command line arguments and convert DOCX to Markdown."""
    parser = argparse.ArgumentParser(description='Convert DOCX files to Markdown format')
    parser.add_argument('input_file', help='Input DOCX file')
    parser.add_argument('-o', '--output', help='Output Markdown file')
    args = parser.parse_args()
    
    try:
        converter = DOCX2Markdown(args.input_file, args.output)
        result = converter.convert()
        
        # If no output file specified, print to stdout
        if not args.output:
            print(result)
            
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
